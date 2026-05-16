#!/usr/bin/env python3
"""
Knowledge Extractor — 安全标准知识抽取管线

从 PDF/TXT 文件中提取安全标准条款，按主题过滤并抽取结构化知识，
输出到 Word 文档。

两种模式:
  mock  — 关键词匹配（无需 LLM，立即跑通 Demo）
  ollama — 本地大模型过滤+抽取（需安装 Ollama）

用法:
  # Demo 模式
  python pipeline.py --input samples/ --output output/demo.docx --topic 液氧 --mode mock

  # Ollama 模式
  python pipeline.py --input samples/ --output output/result.docx --topic 液氧 --mode ollama --model qwen2.5
"""

import os
import re
import json
import argparse
import logging
import sys
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field, asdict
from collections import defaultdict

import requests
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    import fitz
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("extractor")


# ─── 数据结构 ────────────────────────────────────────────────

@dataclass
class KnowledgeItem:
    source_standard: str   # 标准编号/名称
    clause: str            # 条款号
    topic: str             # 主题分类
    requirement: str       # 要求内容
    value: str             # 数值参数
    condition: str         # 适用条件


# ─── 1. 文本提取 ─────────────────────────────────────────────

def extract_text_from_pdf(path: Path) -> str:
    doc = fitz.open(str(path))
    pages = [page.get_text() for page in doc]
    doc.close()
    return "\n".join(pages)


def extract_text_from_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def load_document(path: Path) -> Tuple[str, str]:
    suffix = path.suffix.lower()
    name = path.stem
    if suffix == ".pdf":
        if not HAS_PDF:
            raise RuntimeError("PyMuPDF 未安装: pip install PyMuPDF")
        return extract_text_from_pdf(path), name
    elif suffix == ".txt":
        return extract_text_from_txt(path), name
    else:
        raise ValueError(f"不支持的文件格式: {suffix}，仅支持 .pdf / .txt")


# ─── 2. 按条款分块 ──────────────────────────────────────────

# 匹配常见国标章节号: "4.1" "4.2.1" "5.1.1.2" "第4章" "第5条" "4.1 储存要求"
SECTION_RE = re.compile(
    r'(?:^|\n)(?=\d{1,2}(?:\.\d{1,2}){1,3}\s)|'
    r'(?:^|\n)(?=第[一二三四五六七八九十\d]+[章节条])',
    re.MULTILINE,
)

# 条款行提取：纯数字章节号（4.1、4.2.1、5.1.1.2）
CLAUSE_NUM_RE = re.compile(r'^(\d{1,2}(?:\.\d{1,2}){1,3})')


def chunk_by_sections(text: str) -> List[Dict[str, str]]:
    """将文本按章节标题拆分为多个块，保留章节号。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 找章节边界
    boundaries = [m.start() for m in SECTION_RE.finditer(text)]
    if not boundaries:
        return [{"clause": "全文", "content": text.strip()}] if text.strip() else []

    chunks = []
    for i, start in enumerate(boundaries):
        end = boundaries[i + 1] if i + 1 < len(boundaries) else len(text)
        section = text[start:end].strip()
        if len(section) < 15:
            continue

        # 解析章节号
        first_line = section.split("\n")[0].strip()
        clause_match = CLAUSE_NUM_RE.match(first_line)
        clause = clause_match.group(1) if clause_match else first_line[:25]

        chunks.append({"clause": clause, "content": section})

    return chunks


def split_into_sentences(text: str) -> List[str]:
    """按句号、分号、换行拆分为短句。"""
    parts = re.split(r"[。；\n]", text)
    return [p.strip() for p in parts if len(p.strip()) > 5]


# ─── 3. 主题过滤 ─────────────────────────────────────────────

# 液氧安全关键词
OXYGEN_KEYWORDS = ["液氧", "液氧储罐", "LOX", "loxygen", "低温液体",
                   "深冷", "氧", "氧气储罐", "液氧泵",
                   "充装氧", "liquefied oxygen", "liquid oxygen"]

# 主题分类关键词映射
TOPIC_KEYWORDS: Dict[str, List[str]] = {
    "储运与储存":   ["储罐", "储存", "存放", "贮罐", "储量", "液位", "保冷"],
    "充装与灌装":   ["充装", "灌装", "装车", "装卸", "充灌", "装瓶"],
    "消防与间距":   ["间距", "防火", "消防", "灭火", "安全距离", "隔离", "防火堤"],
    "个人防护":     ["防护", "护具", "面罩", "手套", "防冻", "防护服", "冻伤"],
    "设备与管道":   ["管道", "阀门", "密封", "材料", "真空", "绝热", "材质", "耐压"],
    "操作规程":     ["操作", "严禁", "禁止", "必须", "不得", "不应", "应"],
    "应急处理":     ["应急", "泄漏", "事故", "排放", "堵漏", "疏散", "急救"],
    "运输安全":     ["运输", "槽车", "罐车", "配送", "押运"],
}


def is_relevant_mock(text: str) -> bool:
    """关键词级过滤（Mock 模式）。"""
    lower = text.lower()
    return any(kw in text or kw.lower() in lower for kw in OXYGEN_KEYWORDS)


def is_relevant_ollama(text: str, topic: str, ollama_url: str, model: str) -> bool:
    """LLM 级二分类过滤。"""
    prompt = (
        f"判断以下文本是否与「{topic}」安全相关。只回答「是」或「否」。\n\n"
        f"文本：{text[:1200]}"
    )
    try:
        resp = requests.post(
            f"{ollama_url}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {"num_predict": 10},
            },
            timeout=30,
        )
        resp.raise_for_status()
        ans = resp.json()["response"].strip()
        return "是" in ans or "yes" in ans.lower()
    except Exception as e:
        logger.warning(f"Ollama 过滤调用失败: {e}，回退关键词")
        return is_relevant_mock(text)


def classify_topic(text: str) -> List[str]:
    """通过关键词分类主题（Mock 模式也用）。"""
    matched = []
    for topic, kws in TOPIC_KEYWORDS.items():
        if any(kw in text for kw in kws):
            matched.append(topic)
    return matched if matched else ["通用要求"]


def extract_numeric_value(text: str) -> str:
    """提取句子中的数值参数。"""
    patterns = [
        r"(\d+(?:\.\d+)?\s*(?:m|mm|cm|km|℃|°C|K|%|MPa|kPa|bar|"
        r"L|m³|kg|t|h|min|s|m/s|m³/h))",
        r"(不应[小大]于\s*\d+(?:\.\d+)?)",
        r"(不得[小大]于\s*\d+(?:\.\d+)?)",
        r"([﹤<]\s*\d+(?:\.\d+)?)",
        r"([﹥>]\s*\d+(?:\.\d+)?)",
        r"(-?\d+(?:\.\d+)?\s*℃)",
        r"(\d+(?:\.\d+)?\s*%)",
    ]
    vals = []
    for p in patterns:
        vals.extend(re.findall(p, text))
    return "；".join(v[:20] for v in vals) if vals else ""


# ─── 4. 知识抽取 ─────────────────────────────────────────────

def extract_knowledge_mock(chunk: Dict[str, str], source: str) -> List[KnowledgeItem]:
    """Mock 模式：基于关键词 + 句法规则抽取。"""
    content = chunk["content"]
    clause = chunk["clause"]
    topics = classify_topic(content)
    items: List[KnowledgeItem] = []

    sentences = split_into_sentences(content)
    req_sentences = [
        s for s in sentences
        if any(kw in s for kw in ["应", "不得", "严禁", "必须", "禁止", "不应"])
        and any(kw in s for kw in OXYGEN_KEYWORDS)
    ]

    if not req_sentences:
        # 整块回退
        for t in topics[:2]:
            items.append(KnowledgeItem(
                source_standard=source,
                clause=clause,
                topic=t,
                requirement=content[:300].strip(),
                value=extract_numeric_value(content),
                condition="",
            ))
    else:
        seen = set()
        for s in req_sentences[:6]:
            dedup_key = s[:40]
            if dedup_key in seen:
                continue
            seen.add(dedup_key)
            for t in topics[:2]:
                items.append(KnowledgeItem(
                    source_standard=source,
                    clause=clause,
                    topic=t,
                    requirement=s,
                    value=extract_numeric_value(s),
                    condition="",
                ))

    return items


def extract_knowledge_ollama(
    chunk: Dict[str, str], source: str, topic: str,
    ollama_url: str, model: str,
) -> List[KnowledgeItem]:
    """Ollama 模式：LLM 结构化抽取。"""
    prompt = (
        f"你是一个安全标准专家。从以下文本中提取所有与「{topic}」相关的安全技术要求条款。\n\n"
        f"来源标准：{source}\n"
        f"条款号：{chunk['clause']}\n"
        f"文本：\n{chunk['content'][:2500]}\n\n"
        f"请以 JSON 数组格式输出，每条格式如下，不要输出 JSON 以外的内容：\n"
        f'[\n'
        f'  {{\n'
        f'    "clause": "条款号",\n'
        f'    "topic": "主题分类（储运与储存|充装与灌装|消防与间距|个人防护|设备与管道|操作规程|应急处理|运输安全|通用要求）",\n'
        f'    "requirement": "要求内容",\n'
        f'    "value": "涉及的具体数值（距离、温度、压力等，无则留空）",\n'
        f'    "condition": "适用条件或例外（无则留空）"\n'
        f'  }}\n'
        f']'
    )
    try:
        resp = requests.post(
            f"{ollama_url}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False,
                "options": {"num_predict": 2048, "temperature": 0.1},
            },
            timeout=180,
        )
        resp.raise_for_status()
        raw = resp.json()["response"].strip()

        # 清理可能的 markdown 代码块
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

        json_match = re.search(r"\[.*\]", raw, re.DOTALL)
        if not json_match:
            logger.warning(f"  LLM 输出未包含 JSON: {raw[:100]}")
            return []

        data = json.loads(json_match.group())
        return [
            KnowledgeItem(
                source_standard=source,
                clause=item.get("clause", chunk["clause"]),
                topic=item.get("topic", topic),
                requirement=item.get("requirement", ""),
                value=item.get("value", ""),
                condition=item.get("condition", ""),
            )
            for item in data
            if item.get("requirement", "").strip()
        ]
    except json.JSONDecodeError as e:
        logger.warning(f"  JSON 解析失败: {e}")
        return []
    except Exception as e:
        logger.warning(f"  Ollama 抽取失败: {e}")
        return []


# ─── 5. Word 输出 ─────────────────────────────────────────────

def generate_word(items: List[KnowledgeItem], output_path: str, topic: str):
    """将知识条目输出为结构化 Word 文档。"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── 封面标题 ──
    title = doc.add_heading(f"{topic} — 安全标准知识点合集", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 统计信息 ──
    standards = sorted(set(it.source_standard for it in items))
    topic_counts = defaultdict(int)
    for it in items:
        topic_counts[it.topic] += 1

    doc.add_paragraph(
        f"涉及标准：{len(standards)} 部\n"
        f"主题分类：{len(topic_counts)} 类\n"
        f"知识点总数：{len(items)}"
    )
    doc.add_paragraph(f"标准清单：{'、'.join(standards)}")
    doc.add_paragraph("")  # 间距

    # ── 按主题分节 ──
    groups = defaultdict(list)
    for it in items:
        groups[it.topic].append(it)

    # 条目多的主题排前
    sorted_groups = sorted(groups.items(), key=lambda x: -len(x[1]))

    for topic_name, topic_items in sorted_groups:
        doc.add_heading(f"{topic_name}（{len(topic_items)} 条）", level=1)

        # 创建表格
        cols = 5
        table = doc.add_table(rows=1, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        header_texts = ["标准编号", "条款号", "要求内容", "具体数值", "适用条件"]
        widths = [Cm(3.5), Cm(2.0), Cm(12.0), Cm(3.5), Cm(3.5)]
        for i, (ht, w) in enumerate(zip(header_texts, widths)):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(ht)
            run.bold = True
            run.font.size = Pt(9)
            cell.width = w

        # 数据行
        for item in topic_items:
            row = table.add_row()
            values = [
                item.source_standard,
                item.clause,
                item.requirement,
                item.value,
                item.condition,
            ]
            for i, val in enumerate(values):
                row.cells[i].text = ""
                p = row.cells[i].paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(8.5)
                if widths:
                    row.cells[i].width = widths[i]

        doc.add_paragraph("")  # 节后间距

    doc.save(output_path)
    logger.info(f"Word 文档已保存: {output_path}")


# ─── 6. 主流程 ────────────────────────────────────────────────

def run_pipeline(
    input_dir: str,
    output_path: str,
    topic: str,
    mode: str = "mock",
    ollama_url: str = "http://localhost:11434",
    ollama_model: str = "qwen2.5",
    max_workers: int = 4,
):
    """完整管线：加载 → 分块 → 过滤 → 抽取 → 输出 Word。"""
    input_path = Path(input_dir)
    if not input_path.exists():
        logger.error(f"输入目录不存在: {input_dir}")
        sys.exit(1)

    files = []
    for ext in ("*.pdf", "*.txt"):
        files.extend(sorted(input_path.glob(ext)))
    if not files:
        logger.error(f"输入目录中未找到 .pdf 或 .txt 文件: {input_dir}")
        sys.exit(1)

    logger.info(f"找到 {len(files)} 个文档，模式: {mode}")

    all_items: List[KnowledgeItem] = []

    for fp in files:
        logger.info(f"── 处理: {fp.name} ──")
        try:
            text, source_name = load_document(fp)
            logger.info(f"  文本长度: {len(text)} 字符")
        except Exception as e:
            logger.error(f"  加载失败: {e}")
            continue

        chunks = chunk_by_sections(text)
        logger.info(f"  分块: {len(chunks)} 个章节")

        for chunk in chunks:
            if len(chunk["content"]) < 15:
                continue

            # ── 主题过滤 ──
            if mode == "ollama":
                if not is_relevant_ollama(chunk["content"], topic, ollama_url, ollama_model):
                    continue
            else:
                if not is_relevant_mock(chunk["content"]):
                    continue

            # ── 知识抽取 ──
            if mode == "ollama":
                items = extract_knowledge_ollama(
                    chunk, source_name, topic, ollama_url, ollama_model,
                )
            else:
                items = extract_knowledge_mock(chunk, source_name)

            if items:
                logger.info(f"  ✓ §{chunk['clause']} → {len(items)} 条")
            all_items.extend(items)

    logger.info(f"\n总计: {len(all_items)} 条知识点")

    if not all_items:
        logger.warning("未提取到任何知识点，请检查输入内容或调整关键词。")
        return

    generate_word(all_items, output_path, topic)


# ─── CLI 入口 ─────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="安全标准知识抽取管线 — 从 PDF/TXT 到 Word",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例:\n"
            "  # Mock 模式（无需 LLM）\n"
            "  python pipeline.py -i samples/ -o output/demo.docx -t 液氧\n\n"
            "  # Ollama 模式（需先启动: ollama pull qwen2.5）\n"
            "  python pipeline.py -i samples/ -o output/result.docx -t 液氧 -m ollama\n"
        ),
    )
    parser.add_argument("-i", "--input", default="samples/",
                        help="输入目录（含 .pdf/.txt 文件）")
    parser.add_argument("-o", "--output", default="output/知识合集.docx",
                        help="输出 Word 路径")
    parser.add_argument("-t", "--topic", default="液氧",
                        help="主题关键词，如 液氧、液氢、低温")
    parser.add_argument("-m", "--mode", choices=["mock", "ollama"], default="mock",
                        help="抽取模式: mock（关键词）| ollama（本地 LLM）")
    parser.add_argument("--ollama-url", default="http://localhost:11434",
                        help="Ollama API 地址")
    parser.add_argument("--model", default="qwen2.5",
                        help="Ollama 模型名（如 qwen2.5、llama3.1）")
    parser.add_argument("-w", "--workers", type=int, default=4,
                        help="并发线程数（当前未启用多线程）")
    args = parser.parse_args()

    run_pipeline(
        input_dir=args.input,
        output_path=args.output,
        topic=args.topic,
        mode=args.mode,
        ollama_url=args.ollama_url,
        ollama_model=args.model,
        max_workers=args.workers,
    )


if __name__ == "__main__":
    main()
